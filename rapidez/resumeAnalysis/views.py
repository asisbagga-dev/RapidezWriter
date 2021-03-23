from django.shortcuts import render, redirect, HttpResponse
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from .forms import ResumeForm
from .models import Resume
from django.template.loader import render_to_string
from django.utils.html import strip_tags

# Celery
# from .tasks import sleepy, send_email

# Emails
from django.core.mail import send_mail

# Resume analysis
import docx

# Resume parser
# from pyresparser import ResumeParser

import json


def resume_analysis(request):
    return render(request, 'resume_analysis.html')

def thankyou(request):
    return render(request, 'thankyou.html')

def resume_upload(request):
    if request.method == 'POST':
        form = ResumeForm(request.POST, request.FILES)
        if form.is_valid():
            userResume = form.save()
            # request.session['id'] = form.pk
            
            userResumePath = userResume.resume.path
            print("userResumePath: " + userResumePath)
            resume_parser(userResumePath)
            
            return redirect('resumeAnalysis:thankyou')
    else:
            form = ResumeForm()
    return render(request, 'resume_upload.html', {'form': form})


def resume_parser(resumePath):
    # id = request.session.get('id')
    # print(id)
    # user = Resume.objects.filter(pk=id)
    # for i in user:
    #     resume_file = i.resume.url
    #     print(i.resume.url)
    # doc = docx.Document(resume_file)

    # data = ResumeParser(resumePath).get_extracted_data()
    print("\n")
    pretty_data = json.dumps(data, sort_keys=True, indent=4)
    print(pretty_data)

    doc = docx.Document(resumePath)
    # all_paras = doc.paragraphs

    imagePresentInResume = False
    numberOfPages = 0
    numberOfBullets = 0
    sectionCount = 0
    sections = ['skills', 'education', 'experience', 'objective', 'certification']
    secntionsPresent = []

    for paragraph in doc.paragraphs:
        # If you find an image
        if 'Graphic' in paragraph._p.xml:
            imagePresentInResume = True
        else:
            imagePresentInResume = False
        
        if 'w:lastRenderedPageBreak' in paragraph._p.xml:
            numberOfPages += 1
        
        if 'ListBullet' in paragraph._p.xml:
            numberOfBullets +=1

        for section in sections:
            if section in paragraph._p.xml:
                sectionCount +=1
                secntionsPresent.append(section)

    print("numberOfPages:")
    print(numberOfPages+1)
    if imagePresentInResume == True:
        print("Images or graphic is present in the resume")
    else:
        print("Images or graphic is not present in the resume")

    if numberOfBullets:
        print("Bullets is present in the resume")
        print("numberOfBullets:")
        print(numberOfBullets)
    else:
        print("Bullets is not present in the resume")

    if sectionCount < 2:
        print("Relevant sections are missing")
    elif sectionCount < 4:
        print("Relevant sections are present, maybe can include more optional ones")
    else:
        print(sectionCount)
        print("Relevant sections included")

    # Email 
    email_from = settings.EMAIL_HOST_USER
    email_to = 'bharath.nr1@gmail.com'
    email_admin = settings.EMAIL_ADMIN
    content = 'Resume Analysis'

    html_content = render_to_string('mail_template.html') 
    text_content = strip_tags(html_content)

    # message1 = ('Quote request', 'Have attached the pdf with the required materials', email_from, [email_to])
    # message2 = ('New comedian registration', content , email_from, [email_admin])
    # send_mass_mail((message1, message2), fail_silently=False)
    isSuccess = send_mail(
        'Resume Analysis',
        text_content,
        email_admin,
        [email_to],
        fail_silently=False,
    )
    
    print("Email sent : ", isSuccess)
    return


def test(request):
    print("Inside test function")
    # send_email()

    # doc = docx.Document(r"C:\Stuff\WDF\rapidez\RapidezWriter-1\Samples\Bhattacharya.docx")
    # all_paras = doc.paragraphs
    # print(len(all_paras))
    # return redirect('resumeAnalysis:thankyou')

    # data = ResumeParser(r"C:\Users\bharath878\Downloads\b-suryanarayanan-cv-2020.pdf").get_extracted_data()
    print("\n")
    pretty_data = json.dumps(data, sort_keys=True, indent=4)
    print(pretty_data)
    return render(request, 'thankyou.html', {'pretty_data': pretty_data})